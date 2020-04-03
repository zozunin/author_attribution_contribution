from razdel import tokenize
import docx
from docx import Document
import os
from nltk.tokenize import WordPunctTokenizer
tokenizer = WordPunctTokenizer()
import nltk
nltk.download('stopwords')
from nltk.corpus import stopwords
stop_words = stopwords.words('russian')
import re
import pymorphy2
morph = pymorphy2.MorphAnalyzer()



path = r"C:\Users\user\Desktop\project\docen"
docs = os.listdir(path)

def token_prepro(doc):
    #doc - строкой

    #косые черты
    r1 = re.findall(r'\W\w/\w', doc)
    for case in r1:
        doc = re.sub(case, '_'.join(case.split('/')), doc)
    #нижегородский код
    doc = re.sub('\(831\)', '831', doc)
    # время, деньги
    r3 = re.findall(r'\d+[ :]\d+', doc)
    for case in r3:
        doc = re.sub(case, '.'.join(case.split(' ')), doc)
        doc = re.sub(case, '.'.join(case.split(':')), doc)
    doc = re.sub('\xa0', '', doc)
    doc = re.sub('\?{2,3}', '?', doc)
    doc = re.sub('\!{2,3}', '!', doc)
    doc = re.sub('\!\?', '?', doc)
    doc = re.sub('\?\!', '?', doc)
    return doc

def ww_prepro(doc):
    tokenized = []
    clean_orig = []
    lemmatized = []
    POS_tags = []
    doc = re.sub(r'\*{6}', '', doc)
    for tok in list(tokenize(doc)):
        tokenized.append(tok.text.lower())

    for term in tokenized:
        if morph.parse(term)[0].normal_form not in stop_words and bool(re.match(r'[а-я]+\-?', morph.parse(term)[0].normal_form)) and len(term) > 2:
            tag = morph.parse(term)[0].tag.POS
            if not tag:
                tag = "UNK"
            clean_orig.append(term)
            POS_tags.append(tag)
            lemmatized.append(morph.parse(term)[0].normal_form)
    return [clean_orig,lemmatized,POS_tags]

def ngram_find(n, docs, top=10, sep=False):
    n_gramms = {}
    if sep == False:
        n_gramms['common'] = {}
        for v in docs.values():
            text = ww_prepro(v)
            for term in range(len(text[1])):
                if ' '.join(text[1][term:term+n+1]) not in n_gramms['common']:
                    n_gramms['common'][' '.join(text[1][term:term+n+1])] = 1
                else:
                    n_gramms['common'][' '.join(text[1][term:term+n+1])] += 1
    elif sep == True:
        
        for k, v in docs.items():
            n_gramms[k] = {}
            text = ww_prepro(v)
            for term in range(len(text[1])):
                if ' '.join(text[1][term:term+n]) not in n_gramms[k]:
                    n_gramms[k][' '.join(text[1][term:term+n])] = 1
                else:
                    n_gramms[k][' '.join(text[1][term:term+n])] += 1
    sorted_ngramms = {}
    for k, v in n_gramms.items():
        if len(v) < top:
            top = len(v)
        temp = sorted(v.items(), key=lambda x: x[1], reverse=True)[:top]
        sorted_ngramms[k] = temp
        
    return sorted_ngramms

def write_doc(sorted_ngramms, n):       
    document = Document()
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Файл'
    hdr_cells[1].text = '{}-граммы'.format(n)
    for name, counts in sorted_ngramms.items():
        row_cells = table.add_row().cells
        row_cells[0].text = name
        text = []
        for step in counts:
            if step[1] == 1:
                text = ['Не найдено']
                break
            else:
                text.append(step[0] + ' : ' + str(step[1]))
        row_cells[1].text = '\n'.join(text)
    
    document.add_page_break()
    document.save('N-Gramms.docx') #путь для сохранения, имя

    
in_process = {}

for doc in docs:
    if 'docx' in doc:
        doc_text = []
        temp = docx.Document(path+os.sep+doc)
        for paragraph in temp.paragraphs:
            doc_text.append(paragraph.text)
        for table in temp.tables:
            for row in table.rows:
                for cell in row.cells:
                   doc_text.append(cell.text)

        doc_text = '****** '.join(doc_text)
        
    if 'txt' in doc:
        doc_text = open(path+os.sep+doc, 'r', encoding='utf-8').read()

    in_process[doc] = token_prepro(doc_text)

# sep == True для каждого текста формируется отдельный список n-грамм
# sep ==False (по умолчанию) список n-gramm формируется общий на весь текстовый набор

n = 2 # n-грамма
top = 10 #количество наиболее частотных n-грамм
sorted_ngramms = ngram_find(n, in_process, top, sep=True)
write_doc(sorted_ngramms, n)
