from razdel import tokenize
import docx
from docx import Document
import os
from nltk.tokenize import WordPunctTokenizer
tokenizer = WordPunctTokenizer()
import nltk
import re
import pymorphy2
morph = pymorphy2.MorphAnalyzer()
import numpy as np
from docx import Document


path = r"C:\Users\user\Desktop\project\docen"
docs = os.listdir(path)

def token_prepro(doc):
    #doc - строкой

    #косые черты
    r1 = re.findall(r'\W\w/\w', doc)
    for case in r1:
        doc = re.sub(case, '_'.join(case.split('/')), doc)
    #нижегородский код
    doc = re.sub('\(831\)', '831', doc)
    # время - деньги
    r3 = re.findall(r'\d+[ :]\d+', doc)
    for case in r3:
        doc = re.sub(case, '.'.join(case.split(' ')), doc)
        doc = re.sub(case, '.'.join(case.split(':')), doc)
    doc = re.sub('\xa0', '', doc)
    doc = re.sub('\?{2,3}', '?', doc)
    doc = re.sub('\!{2,3}', '!', doc)
    doc = re.sub('\!\?', '?', doc)
    doc = re.sub('\?\!', '?', doc)
    return doc


def ww_prepro(doc):
    tokenized = []
    lemmatized = []
    POS_tags = []
    doc = re.sub(r'\*{6}', '', doc)
    for tok in list(tokenize(doc)):
        tokenized.append(tok.text.lower())

    for term in tokenized:
        tag = morph.parse(term)[0].tag.POS
        if not tag:
            tag = "UNK"
        POS_tags.append(tag)
        lemmatized.append(morph.parse(term)[0].normal_form)
    return [tokenized,lemmatized,POS_tags]

# разбиение на предложения

short_list = ['др','пр','см', 'п', 'г', 'тел', 'пт', 'ч', 'р', 'ст', 'ср','напр','вв', 'гг','тт','cт','обл','оз','стр','акад','доц','проф','им' 'адм', 'терр','гос','гр','жен',
       'заруб','ин','иностр','инст','муж','мин','миним','просп','пр','руб','стр','ул', 'оф']

def sentence_tokenize(doc):
    messages = doc.split('******')
    split_m = []
    for m in messages:
        split_m.append(m.split('\n'))
	
    done_sentences = []
    for m in split_m:
        temp_s = []
        end = False
        for idx, p in enumerate(m):
            lenm = len(m)
            p = p.strip()
            if not end:
                if bool(re.match(r'^[СC] уважением', p)) and (True in [True for nxt in m[idx+1:lenm] if bool(re.match(r'((8|\+7)[\- ]?)?(\(?\d{3}\)?[\- ]?)?[\d\- ]{7,10}', nxt))]
                                                           or True in [True for nxt in m[idx+1:lenm] if 'Эксперт Союз' in nxt]
                                                           or True in [True for nxt in m[idx+1:lenm] if bool(re.search(r'\d+', nxt))]):
                    temp_s += [' '.join(m[idx:lenm])]
                    end = True
                elif bool(re.match(r'^С уважением', p)):
                    try:
                        if m[idx] == m[-2]:
                            temp_s += [' '.join(m[idx:lenm])]
                            end = True
                        elif m[idx] == m[-1]:
                            temp_s += [p]
                        else:
                            temp_s += [p]
                    except:
                        pass
                else:  
                    temp = nltk.sent_tokenize(p, language='russian')
                    temp_s+=temp
            elif end:
                break
        upd_sentence = []
        included = []
        for idx, p in enumerate(temp_s):
            if True in [True for w in short_list if bool(re.search(r' {}.$'.format(w), p))]:
                try:
                    if not temp_s[idx+1][0].isupper():
                        upd_sentence.append(' '.join(temp_s[idx:idx+2]))
                        included += [idx,idx+1]
                    else:
                        upd_sentence.append(p)
                except:
                    pass
            elif re.findall(r'^\d+\.$', p):
                if p == re.findall(r'^\d+\.$', p)[0]:
                    upd_sentence.append(' '.join(temp_s[idx:idx+2]))
                    included += [idx,idx+1]
            elif idx not in included:
                upd_sentence.append(p)               
                    
        if upd_sentence != []:
            done_sentences+=upd_sentence

    return done_sentences



in_process = {}


for doc in docs:
    if 'docx' in doc:
        doc_text = []
        temp = docx.Document(path+os.sep+doc)
        for paragraph in temp.paragraphs:
            doc_text.append(paragraph.text)
        for table in temp.tables:
            for row in table.rows:
                for cell in row.cells:
                   doc_text.append(cell.text)

        doc_text = '******'.join(doc_text)
        
    if 'txt' in doc:
        doc_text = open(path+os.sep+doc, 'r', encoding='utf-8').read()

    in_process[doc] = token_prepro(doc_text)

sentences = {}
for k, v in in_process.items():
    temp = ''
    if 'docx' in k:
        temp = sentence_tokenize(v)
    elif 'txt' in k:
        temp = nltk.sent_tokenize(v, language='russian')
    full_tok = []
    for part in temp:
        part_tok = []
        for tok in list(tokenize(part)):
            if re.match(r'[\w]', tok.text):
                part_tok.append(tok.text.lower())
        full_tok.append(part_tok)
        
    sentences[k] = full_tok

def write_doc(statistics):
    document = Document()
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Имя документа'
    hdr_cells[1].text = 'Части речи'
    hdr_cells[2].text = 'Статистика по тексту'
    for k, v in statistics.items():
        row_cells = table.add_row().cells
        row_cells[0].text = k
        temp_pos = [str(pos)+' : '+str(c) for pos, c in v['pos_freq'].items() if str(pos)!='UNK']
        row_cells[1].text = '\n'.join(temp_pos)
        row_cells[2].text = '\n'.join(v['sentence_statistics'])
    document.add_page_break()
    document.save('Statistics.docx') # путь сохранения


statistics = {}
for name, text in in_process.items():
    doc = ww_prepro(text)
    statistics[name] = {}
    pos_dict = {}
    for pos in doc[2]:
        if pos not in pos_dict:
            pos_dict[pos] = 1
        else:
            pos_dict[pos] += 1
                
    statistics[name]['pos_freq'] = pos_dict
    statistics[name]['sentence_statistics'] = ['Количество предложений: {}'.format(str(len(sentences[name])))]
    sentence_lengths = []
    word_lengths = []
    for idx, sent_text in enumerate(sentences[name]):
        sentence_lengths.append(len(sent_text))
        for word in sent_text:
            word_lengths.append(len(word))
    statistics[name]['sentence_statistics'] += ['Средняя длина предложний: {}'.format(str(round(np.mean(sentence_lengths), 2)))]
    statistics[name]['sentence_statistics'] += ['Средняя длина слов: {}'.format(str(round(np.mean(word_lengths), 2)))]

write_doc(statistics)
