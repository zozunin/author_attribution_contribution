from razdel import tokenize
import docx
from docx import Document
import os
from nltk.tokenize import WordPunctTokenizer
tokenizer = WordPunctTokenizer()
import nltk
import re
import pymorphy2
morph = pymorphy2.MorphAnalyzer()


path = r"C:\Users\user\Desktop\project\docen"# путь до папки с файлами формата txt\docx
docs = os.listdir(path)

def token_prepro(doc):
    #doc - строкой

    #косые черты
    r1 = re.findall(r'\W\w/\w', doc)
    for case in r1:
        doc = re.sub(case, '_'.join(case.split('/')), doc)
    #нижегородский код
    doc = re.sub('\(831\)', '831', doc)
    # время - деньги
    r3 = re.findall(r'\d+[ :]\d+', doc)
    for case in r3:
        doc = re.sub(case, '.'.join(case.split(' ')), doc)
        doc = re.sub(case, '.'.join(case.split(':')), doc)
    doc = re.sub('\xa0', '', doc)
    doc = re.sub('\?{2,3}', '?', doc)
    doc = re.sub('\!{2,3}', '!', doc)
    doc = re.sub('\!\?', '?', doc)
    doc = re.sub('\?\!', '?', doc)
    return doc

def ww_prepro(doc):
    tokenized = []
    lemmatized = []
    POS_tags = []
    doc = re.sub(r'\*{6}', ' ', doc)
    for tok in list(tokenize(doc)):
        tokenized.append(tok.text.lower())

    for term in tokenized:
        tag = str(morph.parse(term)[0].tag.POS)
        if not tag:
            tag = "UNK"
        POS_tags.append(tag)
        lemmatized.append(morph.parse(term)[0].normal_form)
    return [tokenized,lemmatized,POS_tags]

def get_intensifiers():
    intense_list = ''
    temp = docx.Document(r"C:\Users\user\Desktop\project\Spisok_intensifikatorov.docx") # путь до списка
    for table in temp.tables:
        for row in table.rows:
            for cell in row.cells:
                if len(cell.text) > 1:
                    intense_list += cell.text + '\n'
    clean_list = []
    for word in intense_list.split('\n')[:-1]:
        if '(' in  word or ')' in word:
            temp = re.sub(r'[()]', '', word)
            temp = temp.split()
            clean_list.append(temp)
            #clean_list.append(re.sub(r'[ ]?\(.+\)[ ]?', '', word))
            #clean_list.append(re.sub(r'[()]', '', word))
        else:
            clean_list.append(word)
    return clean_list

def save_context(doc, n, tok):
    sample = doc[0][tok-n:tok]+[doc[0][tok].upper()]+doc[0][tok+1:tok+n]
    return ' '.join(sample)

def find_intensifiers(doc, intensifiers, n=3):
    context_intense = {}
    for word in intensifiers:
        values = []
        for tok in range(len(doc[1])):
            if type(word) is list: #случаи как вовсе не
                if doc[1][tok:tok+2] == word:
                    if doc[2][tok+2] in ['ADJF', 'ADJS', 'COMP', 'PRTF', 'PRTS', 'ADVB']:
                        values.append(save_context(doc, n, tok))
                            
                elif doc[1][tok] == word[0] and doc[1][tok+1] != word[1]:
                    if doc[2][tok+1] in ['ADJF', 'ADJS', 'COMP', 'PRTF', 'PRTS', 'ADVB']:
                        values.append(save_context(doc, n, tok))
            else:
                if doc[1][tok] == word:
                    if word == 'действительно' and (',' in doc[1][tok-1] or ',' in doc[2][tok+1]):
                        values.append(save_context(doc, n, tok))
                    elif word == 'какой':
                        if 'ADJ' in doc[2][tok-1] and 'ADJ' in doc[2][tok+1]:
                            pass
                        elif 'ADJ' not in doc[2][tok-1] and 'ADJ' in doc[2][tok+1]:
                            values.append(save_context(doc, n, tok))
                        elif doc[2][tok+1] in ['NOUN', 'NPRO']:
                            values.append(save_context(doc, n, tok))
                        elif doc[2][tok+1] == 'ADVB' and doc[2][tok+2] in ['ADJF', 'ADJS', 'NOUN', 'NPRO']:
                            values.append(save_context(doc, n, tok))
                    elif word in ['настоящий', 'страшный', 'невероятный', 'сущий', 'чистый'] and doc[2][tok+1] == 'NOUN':
                        values.append(save_context(doc, n, tok))
                    elif word in ['немного', 'страшно'] and 'ADJ' in doc[2][tok+1]:
                        values.append(save_context(doc, n, tok))
                    elif word == 'чуть':
                        if doc[1][tok+1] in ['ли', 'не'] or 'ADJ' in doc[2][tok+1]:
                            values.append(save_context(doc, n, tok))
                    elif word == 'целиком' and doc[1][tok+2] == 'полностью':
                        values.append(save_context(doc, n, tok))
                    elif word == 'так':
                        if doc[1][tok+1] in ['сказать','далее'] or ('называем' in doc[0][tok+1] and 'PRT' in doc[0][tok+2]):
                            pass
                        else:
                            if doc[2][tok+1] in ['ADJS', 'PRTF', 'PRTS', 'ADVB']:
                                values.append(save_context(doc, n, tok))
                            elif doc[2][tok-1] in ['ADJS','ADVB']:
                                values.append(save_context(doc, n, tok))
                    elif word == 'такой':
                        if doc[2][tok+1] in ['ADJF', 'NPRO', 'NOUN']:
                            values.append(save_context(doc, n, tok))
                        elif doc[2][tok+1] == 'ADVB' and doc[2][tok+1] in ['ADJF', 'NOUN', 'NPRO']:
                            values.append(save_context(doc, n, tok))

                             
                    elif word == 'много':
                        pass #возможно уточнение
                    else:
                        if doc[2][tok+1] in ['ADJF', 'ADJS', 'COMP', 'PRTF', 'PRTS', 'ADVB']:
                                values.append(save_context(doc, n, tok))

        if type(word) is list:
            word = '+'.join(word)
        context_intense[word] = values
    return context_intense

def write_doc(name, doc_intensifiers):
    document = Document()
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Интенсификатор'
    hdr_cells[1].text = 'Контексты'
    hdr_cells[2].text = 'Абсолютная частота'
    for k, v in doc_intensifiers.items():
        if v != []:
            row_cells = table.add_row().cells
            row_cells[0].text = k
            row_cells[1].text = '\n\n'.join(v)
            row_cells[2].text = str(len(v))
    document.add_page_break()
    if 'txt' in name:
        name = name[:-3]+'docx'
    document.save('pos_intensifiers_'+name) #+путь до папки сохранения

#MAIN
in_process = {}


for doc in docs:
    if 'docx' in doc:
        doc_text = []
        temp = docx.Document(path+os.sep+doc)
        for paragraph in temp.paragraphs:
            doc_text.append(paragraph.text)
        for table in temp.tables:
            for row in table.rows:
                for cell in row.cells:
                   doc_text.append(cell.text)

        doc_text = '******'.join(doc_text)
        
    if 'txt' in doc:
        doc_text = open(path+os.sep+doc, 'r', encoding='utf-8').read()

    in_process[doc] = token_prepro(doc_text)

intensifiers = get_intensifiers()

for name, text in in_process.items():
    doc = ww_prepro(text)
    doc_intensifiers = find_intensifiers(doc, intensifiers)
    write_doc(name, doc_intensifiers)

